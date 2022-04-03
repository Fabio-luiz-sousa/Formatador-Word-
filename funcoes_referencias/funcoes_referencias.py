import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE

doc = Document()
styles = doc.styles


def titulo_ref():
    estilo_titulo_ref = styles.add_style(
        'Estilo_Titulo_Ref', WD_STYLE_TYPE.PARAGRAPH)
    estilo_titulo_ref.font.name = 'Arial'
    estilo_titulo_ref.font.size = Pt(14)
    estilo_titulo_ref.font.bold = True
    p_titulo_ref = doc.add_paragraph(
        'Referências Bibliográficas', 'Estilo_Titulo_Ref')
    p_titulo_ref_format = p_titulo_ref.paragraph_format
    p_titulo_ref_format.space_after = Pt(30)
    p_titulo_ref.alignment = 1


def ref_livro(nome_autor, sobrenome_autor, titulo, edicao, local_publicacao, editora, data_publicacao):
    estilo_ref_livro = styles.add_style(
        'Estilo_Ref_Livro', WD_STYLE_TYPE.PARAGRAPH)
    estilo_ref_livro.font.name = 'Arial'
    estilo_ref_livro.font.size = Pt(12)

    if len(nome_autor)==1:
        p_ref_livro = doc.add_paragraph(
            f'{sobrenome_autor}, {nome_autor}. {titulo}. {edicao}. {local_publicacao}: {editora}, {data_publicacao}', 'Estilo_Ref_Livro')
        p_ref_livro.alignment = 3
    elif len(nome_autor)<3:
        nome_autor[0]+';'
        for n in nome_autor: # arrumar aqui
            n=nome_autor
        for s in sobrenome_autor:
            s=sobrenome_autor
        p_ref_livro = doc.add_paragraph(
            f'{s}, {n}. {titulo}. {edicao}. {local_publicacao}: {editora}, {data_publicacao}', 'Estilo_Ref_Livro')
        p_ref_livro.alignment = 3
    else:
        p_ref_livro = doc.add_paragraph(
            f'{sobrenome_autor[0]}, {nome_autor[0]}. et al. {titulo}. {edicao}. {local_publicacao}: {editora}, {data_publicacao}', 'Estilo_Ref_Livro')
        p_ref_livro.alignment = 3




    

n=['Fábio','Zeze']
s=['LUIZ','MIRANDA']
titulo_ref()

ref_livro(n, s, 'Historia concisa da literatura brasileira',
          '38.ed', 'São Paulo', 'Cultrix', '1994')

doc.save('tt.docx')