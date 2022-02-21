import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE


doc = Document()


def margens_pagina():
    sections = doc.sections
    for section in sections:  # Coloca as margens da folha no padrao ABNT
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)


margens_pagina()

styles = doc.styles


def titulo(titulo):
    estilo_titulo = styles.add_style('Estilo_Titulo', WD_STYLE_TYPE.PARAGRAPH)
    estilo_titulo.font.name = 'Arial'
    estilo_titulo.font.size = Pt(14)
    estilo_titulo.font.bold = True
    p_titulo = doc.add_paragraph(f'{titulo}', 'Estilo_Titulo')
    p_titulo_format = p_titulo.paragraph_format
    p_titulo_format.space_after = Pt(0)
    p_titulo.alignment = 1


def tipo_curso(curso):
    estilo_tipo_curso = styles.add_style(
        'Estilo_Tipo_Curso', WD_STYLE_TYPE.PARAGRAPH)
    estilo_tipo_curso.font.name = 'Arial'
    estilo_tipo_curso.font.size = Pt(14)
    estilo_tipo_curso.font.bold = True
    p_tipo_curso = doc.add_paragraph(f'{curso}', 'Estilo_Tipo_Curso')
    p_tipo_curso_format = p_tipo_curso.paragraph_format
    p_tipo_curso_format.space_after = Pt(50)
    p_tipo_curso.alignment = 1


def nome_autores(nome_autores):
    nome_autores = sorted(nome_autores)
    estilo_nome_autores = styles.add_style(
        'Estilo_Nome_Autores', WD_STYLE_TYPE.PARAGRAPH)
    estilo_nome_autores.font.name = 'Arial'
    estilo_nome_autores.font.size = Pt(12)
    for n in nome_autores:
        p_nome_autores = doc.add_paragraph(
            f'{n}', 'Estilo_Nome_Autores')
        p_nome_autores_format = p_nome_autores.paragraph_format
        p_nome_autores.alignment = 1
        p_nome_autores_format.space_after = Pt(0)
    c = 8
    if len(nome_autores) > 1:
        c = c-len(nome_autores)*0.3
    p_nome_autores_format.space_after = Cm(c)


def titulo_trabalho(titulo_trabalho):
    estilo_titulo_trabalho = styles.add_style(
        'Estilo_Titulo_Trabalho', WD_STYLE_TYPE.PARAGRAPH)
    estilo_titulo_trabalho.font.name = 'Arial'
    estilo_titulo_trabalho.font.size = Pt(14)
    estilo_titulo_trabalho.font.bold = True
    p_titulo_trabalho = doc.add_paragraph(
        f'{titulo_trabalho}', 'Estilo_Titulo_Trabalho')
    p_titulo_trabalho_format = p_titulo_trabalho.paragraph_format
    p_titulo_trabalho_format.space_after = Pt(0)
    p_titulo_trabalho.alignment = 1


def subtitulo_trabalho(subtitulo_trabalho):
    estilo_subtitulo_trabalho = styles.add_style(
        'Estilo_Subtitulo_Trabalho', WD_STYLE_TYPE.PARAGRAPH)
    estilo_subtitulo_trabalho.font.name = 'Arial'
    estilo_subtitulo_trabalho.font.size = Pt(12)
    p_subtitulo_trabalho = doc.add_paragraph(
        f'{subtitulo_trabalho}', 'Estilo_Subtitulo_Trabalho')
    p_subtitulo_trabalho.alignment = 1
    p_subtitulo_trabalho_format = p_subtitulo_trabalho.paragraph_format
    p_subtitulo_trabalho_format.space_after = Cm(8.3)


def cidade(cidade):
    estilo_cidade = styles.add_style('Estilo_Cidade', WD_STYLE_TYPE.PARAGRAPH)
    estilo_cidade.font.name = 'Arial'
    estilo_cidade.font.size = Pt(12)
    p_cidade = doc.add_paragraph(f'\n\n{cidade}', 'Estilo_Cidade')
    p_cidade.alignment = 1
    p_cidade_format = p_cidade.paragraph_format
    p_cidade_format.space_after = Pt(0)


def ano(ano):
    estilo_ano = styles.add_style('Estilo_Ano', WD_STYLE_TYPE.PARAGRAPH)
    estilo_ano.font.name = 'Arial'
    estilo_ano.font.size = Pt(12)
    p_ano = doc.add_paragraph(f'{ano}', 'Estilo_Ano')
    p_ano_format = p_ano.paragraph_format
    p_ano_format.space_after = Pt(25)
    p_ano.alignment = 1


def nome_autores_contracapa(nome_autores):
    nome_autores = sorted(nome_autores)
    estilo_nome_autores = styles.add_style(
        'Estilo_Nome_Autores_Contra', WD_STYLE_TYPE.PARAGRAPH)
    estilo_nome_autores.font.name = 'Arial'
    estilo_nome_autores.font.size = Pt(12)
    for n in nome_autores:
        p_nome_autores = doc.add_paragraph(
            f'{n}', 'Estilo_Nome_Autores_Contra')
        p_nome_autores_format = p_nome_autores.paragraph_format
        p_nome_autores.alignment = 1
        p_nome_autores_format.space_after = Pt(0)
    c = 11
    if len(nome_autores) > 1:
        c = c-len(nome_autores)*0.3
    p_nome_autores_format.space_after = Cm(c)


def titulo_trabalho_contracapa(titulo_trabalho):
    estilo_titulo_trabalho = styles.add_style(
        'Estilo_Titulo_Trabalho_Contra', WD_STYLE_TYPE.PARAGRAPH)
    estilo_titulo_trabalho.font.name = 'Arial'
    estilo_titulo_trabalho.font.size = Pt(14)
    estilo_titulo_trabalho.font.bold = True
    p_titulo_trabalho = doc.add_paragraph(
        f'{titulo_trabalho}', 'Estilo_Titulo_Trabalho_Contra')
    p_titulo_trabalho_format = p_titulo_trabalho.paragraph_format
    p_titulo_trabalho_format.space_after = Pt(0)
    p_titulo_trabalho.alignment = 1


def subtitulo_trabalho_contracapa(subtitulo_trabalho):
    estilo_subtitulo_trabalho = styles.add_style(
        'Estilo_Subtitulo_Trabalho_Contra', WD_STYLE_TYPE.PARAGRAPH)
    estilo_subtitulo_trabalho.font.name = 'Arial'
    estilo_subtitulo_trabalho.font.size = Pt(12)
    p_subtitulo_trabalho = doc.add_paragraph(
        f'{subtitulo_trabalho}', 'Estilo_Subtitulo_Trabalho_Contra')
    p_subtitulo_trabalho.alignment = 1
    p_subtitulo_trabalho_format = p_subtitulo_trabalho.paragraph_format
    p_subtitulo_trabalho_format.space_after = Cm(2)


def descricao_contracapa(curso, universidade, orientador):
    estilo_descricao_contracapa = styles.add_style(
        'Estilo_Descricao_Contra', WD_STYLE_TYPE.PARAGRAPH)
    estilo_descricao_contracapa.font.name = 'Arial'
    estilo_descricao_contracapa.font.size = Pt(12)
    p_descricao_contracapa = doc.add_paragraph(
        f'Trabalho Apresentado no curso de\n{curso} da {universidade}\n\nOrientador: {orientador}', 'Estilo_Descricao_Contra')
    p_descricao_contracapa.alignment = 0
    p_descricao_contracapa_format = p_descricao_contracapa.paragraph_format
    p_descricao_contracapa_format.left_indent=Cm(8)
    p_descricao_contracapa_format.space_after = Cm(4)


def cidade_contracapa(cidade):
    estilo_cidade = styles.add_style(
        'Estilo_Cidade_Contra', WD_STYLE_TYPE.PARAGRAPH)
    estilo_cidade.font.name = 'Arial'
    estilo_cidade.font.size = Pt(12)
    p_cidade = doc.add_paragraph(f'\n\n{cidade}', 'Estilo_Cidade_Contra')
    p_cidade.alignment = 1
    p_cidade_format = p_cidade.paragraph_format
    p_cidade_format.space_after = Pt(0)


def ano_contracapa(ano):
    estilo_ano = styles.add_style('Estilo_Ano_Contra', WD_STYLE_TYPE.PARAGRAPH)
    estilo_ano.font.name = 'Arial'
    estilo_ano.font.size = Pt(12)
    p_ano = doc.add_paragraph(f'{ano}', 'Estilo_Ano_Contra')
    p_ano_format = p_ano.paragraph_format
    p_ano_format.space_after = Pt(20)
    p_ano.alignment = 1


def salvar_doc(nome_documento):
    doc.save(os.path.abspath('formatador_word')+rf'\{nome_documento}.docx')
